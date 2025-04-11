from base64 import b64encode
import os
import uuid
import time
import json
import threading
import logging
from logging.handlers import RotatingFileHandler
import traceback
from flask import Flask, render_template, request, jsonify, send_from_directory, abort, Response
import openai
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
# Create logs directory if it doesn't exist
LOGS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs')
os.makedirs(LOGS_DIR, exist_ok=True)

# Set up file handler with rotation
log_file = os.path.join(LOGS_DIR, 'app.log')
file_handler = RotatingFileHandler(log_file, maxBytes=10485760, backupCount=10)
file_handler.setFormatter(logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s'))
file_handler.setLevel(logging.INFO)

# Set up console handler
console_handler = logging.StreamHandler()
console_handler.setFormatter(logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s'))
console_handler.setLevel(logging.INFO)

# Configure root logger
logging.basicConfig(level=logging.INFO, handlers=[file_handler, console_handler])
logger = logging.getLogger(__name__)

logger.info("Starting Latin Processing Web Application")

app = Flask(__name__)

# Check if running on Render
IS_RENDER = os.environ.get('RENDER') == 'true'
logger.info(f"Running on Render: {IS_RENDER}")

# Configuration for storage paths
if IS_RENDER:
    # Use Render persistent disk if available
    PERSISTENT_DIR = os.environ.get('RENDER_PERSISTENT_DIR', '/var/data')
    logger.info(f"Using Render persistent directory: {PERSISTENT_DIR}")
    
    # Create persistent directories
    UPLOAD_FOLDER = os.path.join(PERSISTENT_DIR, 'uploads')
    PROCESSED_FOLDER = os.path.join(PERSISTENT_DIR, 'processed')
else:
    # Local development paths
    UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
    PROCESSED_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'processed')

logger.info(f"Upload folder: {UPLOAD_FOLDER}")
logger.info(f"Processed folder: {PROCESSED_FOLDER}")

ALLOWED_EXTENSIONS = {'docx'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16 MB max upload size

# Create directories if they don't exist
try:
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    logger.info(f"Created or verified upload folder: {UPLOAD_FOLDER}")
except Exception as e:
    logger.error(f"Error creating upload folder: {str(e)}")
    logger.error(traceback.format_exc())

try:
    os.makedirs(PROCESSED_FOLDER, exist_ok=True)
    logger.info(f"Created or verified processed folder: {PROCESSED_FOLDER}")
except Exception as e:
    logger.error(f"Error creating processed folder: {str(e)}")
    logger.error(traceback.format_exc())

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['PROCESSED_FOLDER'] = PROCESSED_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Task storage
tasks = {}

# Helper functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Latin correction prompt template
LATIN_CORRECTION_PROMPT = """
You are an expert in early 16th century Latin manuscripts. Your task is to correct transcription errors in the following Latin text while staying very close to the original.

Guidelines:
1. Focus only on fixing obvious transcription errors
2. Preserve period-specific abbreviations and spelling characteristics
3. Make minimal changes to the text
4. Do not modernize or standardize the Latin
5. Preserve the original style and tone

Original Latin text:
{latin_text}

Provide only the corrected Latin text without any explanations or comments:
"""

# Dutch translation prompt template
DUTCH_TRANSLATION_PROMPT = """
You are an expert translator of early 16th century Latin to modern Dutch. Translate the following Latin text into convivial, accessible Dutch.

Guidelines:
1. Create natural, conversational Dutch that modern readers can easily understand
2. Maintain fidelity to the original Latin meaning and tone
3. Preserve the warmth and personality of the original correspondence
4. Use accessible language while respecting the historical context

Latin text to translate:
{latin_text}

Provide only the Dutch translation without any explanations or comments:
"""

def correct_latin_with_chatgpt(text):
    """Correct Latin text using ChatGPT"""
    try:
        logger.info("Starting Latin correction")
        # Check if OPENAI_API_KEY is set
        api_key = os.environ.get('OPENAI_API_KEY')
        if not api_key:
            logger.warning("OPENAI_API_KEY not set, using placeholder correction")
            return text + " [CORRECTED]"
        
        logger.info("OpenAI API key is set")
        openai.api_key = api_key
        
        # Split text into manageable chunks (4000 characters)
        chunks = []
        chunk_size = 4000
        for i in range(0, len(text), chunk_size):
            chunks.append(text[i:i+chunk_size])
        
        logger.info(f"Split text into {len(chunks)} chunks for processing")
        
        corrected_chunks = []
        for i, chunk in enumerate(chunks):
            logger.info(f"Processing chunk {i+1}/{len(chunks)}")
            # Prepare the prompt
            prompt = LATIN_CORRECTION_PROMPT.format(latin_text=chunk)
            
            # Make API call with retry logic
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    logger.info(f"Making OpenAI API call for Latin correction (attempt {attempt+1}/{max_retries})")
                    response = openai.ChatCompletion.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": "You are an expert in early 16th century Latin manuscripts."},
                            {"role": "user", "content": prompt}
                        ],
                        temperature=0.3,
                        max_tokens=4000,
                        timeout=30
                    )
                    
                    corrected_text = response.choices[0].message.content.strip()
                    logger.info(f"Successfully received corrected text for chunk {i+1}")
                    corrected_chunks.append(corrected_text)
                    break
                except Exception as e:
                    logger.error(f"Error in ChatGPT API call (attempt {attempt+1}/{max_retries}): {str(e)}")
                    logger.error(traceback.format_exc())
                    if attempt == max_retries - 1:
                        logger.warning(f"All retries failed for chunk {i+1}, using original text")
                        corrected_chunks.append(chunk)
                    else:
                        time.sleep(2 ** attempt)  # Exponential backoff
        
        logger.info("Latin correction completed successfully")
        return "\n".join(corrected_chunks)
    except Exception as e:
        logger.error(f"Error in Latin correction: {str(e)}")
        logger.error(traceback.format_exc())
        return text + " [ERROR IN CORRECTION]"

def translate_latin_to_dutch_with_chatgpt(text):
    """Translate Latin text to Dutch using ChatGPT"""
    try:
        logger.info("Starting Dutch translation")
        # Check if OPENAI_API_KEY is set
        api_key = os.environ.get('OPENAI_API_KEY')
        if not api_key:
            logger.warning("OPENAI_API_KEY not set, using placeholder translation")
            return "[DUTCH TRANSLATION PLACEHOLDER]"
        
        logger.info("OpenAI API key is set")
        openai.api_key = api_key
        
        # Split text into manageable chunks (3000 characters)
        chunks = []
        chunk_size = 3000
        for i in range(0, len(text), chunk_size):
            chunks.append(text[i:i+chunk_size])
        
        logger.info(f"Split text into {len(chunks)} chunks for translation")
        
        translated_chunks = []
        for i, chunk in enumerate(chunks):
            logger.info(f"Translating chunk {i+1}/{len(chunks)}")
            # Prepare the prompt
            prompt = DUTCH_TRANSLATION_PROMPT.format(latin_text=chunk)
            
            # Make API call with retry logic
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    logger.info(f"Making OpenAI API call for Dutch translation (attempt {attempt+1}/{max_retries})")
                    response = openai.ChatCompletion.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": "You are an expert translator of early 16th century Latin to modern Dutch."},
                            {"role": "user", "content": prompt}
                        ],
                        temperature=0.4,
                        max_tokens=4000,
                        timeout=30
                    )
                    
                    translated_text = response.choices[0].message.content.strip()
                    logger.info(f"Successfully received translation for chunk {i+1}")
                    translated_chunks.append(translated_text)
                    break
                except Exception as e:
                    logger.error(f"Error in ChatGPT API call (attempt {attempt+1}/{max_retries}): {str(e)}")
                    logger.error(traceback.format_exc())
                    if attempt == max_retries - 1:
                        logger.warning(f"All retries failed for chunk {i+1}, using placeholder")
                        translated_chunks.append(f"[TRANSLATION ERROR FOR: {chunk[:100]}...]")
                    else:
                        time.sleep(2 ** attempt)  # Exponential backoff
        
        logger.info("Dutch translation completed successfully")
        return "\n".join(translated_chunks)
    except Exception as e:
        logger.error(f"Error in Dutch translation: {str(e)}")
        logger.error(traceback.format_exc())
        return "[ERROR IN TRANSLATION]"

def create_three_column_document(corrected_latin, dutch_translation, output_path):
    """Create a document with three columns (Latin, spacing, Dutch)"""
    try:
        logger.info(f"Creating three-column document at {output_path}")
        doc = Document()
        
        # Set A3 landscape orientation
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(42.0)  # A3 width
        section.page_height = Cm(29.7)  # A3 height
        
        # Set margins
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        
        # Add title
        title = doc.add_paragraph()
        title_run = title.add_run("Latin Text and Dutch Translation")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle with date
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(f"Processed on {time.strftime('%Y-%m-%d')}")
        subtitle_run.font.size = Pt(12)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # Create table with three columns
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Set column widths
        # First column (Latin): 40% of available width
        # Middle column (spacing): 20% of available width
        # Third column (Dutch): 40% of available width
        table.autofit = False
        table.allow_autofit = False
        
        # Calculate available width (A3 width minus margins)
        available_width = section.page_width - section.left_margin - section.right_margin
        
        # Set column widths
        table.columns[0].width = int(available_width * 0.4)
        table.columns[1].width = int(available_width * 0.2)
        table.columns[2].width = int(available_width * 0.4)
        
        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Latin Text"
        header_cells[1].text = ""  # Empty middle column
        header_cells[2].text = "Dutch Translation"
        
        # Style headers
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(14)
        
        # Split text into paragraphs
        latin_paragraphs = corrected_latin.split('\n')
        dutch_paragraphs = dutch_translation.split('\n')
        
        # Ensure both lists have the same length
        max_paragraphs = max(len(latin_paragraphs), len(dutch_paragraphs))
        latin_paragraphs = latin_paragraphs + [''] * (max_paragraphs - len(latin_paragraphs))
        dutch_paragraphs = dutch_paragraphs + [''] * (max_paragraphs - len(dutch_paragraphs))
        
        # Add content rows
        for latin_para, dutch_para in zip(latin_paragraphs, dutch_paragraphs):
            # Skip empty paragraphs
            if not latin_para.strip() and not dutch_para.strip():
                continue
                
            row = table.add_row()
            cells = row.cells
            
            # Add Latin text
            cells[0].text = latin_para
            
            # Middle column remains empty
            cells[1].text = ""
            
            # Add Dutch translation
            cells[2].text = dutch_para
            
            # Style text
            for cell in cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
        
        # Remove borders from table
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(12)
                cell.border.top = cell.border.bottom = cell.border.left = cell.border.right = None
        
        # Save document
        logger.info(f"Saving document to {output_path}")
        try:
            doc.save(output_path)
            logger.info(f"Document saved successfully to {output_path}")
            
            # Verify file exists after saving
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                logger.info(f"Verified file exists at {output_path} with size {file_size} bytes")
            else:
                logger.error(f"File does not exist after saving: {output_path}")
                
            return True
        except Exception as e:
            logger.error(f"Error saving document to {output_path}: {str(e)}")
            logger.error(traceback.format_exc())
            return False
    except Exception as e:
        logger.error(f"Error creating document: {str(e)}")
        logger.error(traceback.format_exc())
        return False

def compile_documents(processed_files, output_path):
    """Compile all processed documents into a single document"""
    try:
        logger.info(f"Compiling documents into {output_path}")
        logger.info(f"Number of files to compile: {len(processed_files)}")
        for i, file_path in enumerate(processed_files):
            logger.info(f"File {i+1}: {file_path}")
            if not os.path.exists(file_path):
                logger.error(f"File does not exist: {file_path}")
            else:
                file_size = os.path.getsize(file_path)
                logger.info(f"File exists with size {file_size} bytes")
        
        compiled_doc = Document()
        
        # Set A3 landscape orientation
        section = compiled_doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(42.0)  # A3 width
        section.page_height = Cm(29.7)  # A3 height
        
        # Set margins
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        
        # Add title
        title = compiled_doc.add_paragraph()
        title_run = title.add_run("Compiled Latin Texts and Dutch Translations")
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle with date
        subtitle = compiled_doc.add_paragraph()
        subtitle_run = subtitle.add_run(f"Compiled on {time.strftime('%Y-%m-%d')}")
        subtitle_run.font.size = Pt(14)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add table of contents title
        compiled_doc.add_paragraph()
        toc_title = compiled_doc.add_paragraph()
        toc_title_run = toc_title.add_run("Table of Contents")
        toc_title_run.font.size = Pt(16)
        toc_title_run.font.bold = True
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add table of contents
        toc = compiled_doc.add_paragraph()
        
        # Process each file
        for i, file_path in enumerate(processed_files):
            # Get filename without extension
            filename = os.path.basename(file_path)
            name_without_ext = os.path.splitext(filename)[0]
            
            # Add to table of contents
            toc_entry = toc.add_run(f"{i+1}. {name_without_ext}\n")
            toc_entry.font.size = Pt(12)
            
            # Add page break before each document (except the first one)
            if i > 0:
                compiled_doc.add_page_break()
            else:
                compiled_doc.add_paragraph()
                compiled_doc.add_paragraph()
            
            # Add document title
            doc_title = compiled_doc.add_paragraph()
            doc_title_run = doc_title.add_run(f"{i+1}. {name_without_ext}")
            doc_title_run.font.size = Pt(16)
            doc_title_run.font.bold = True
            doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add spacing
            compiled_doc.add_paragraph()
            
            try:
                logger.info(f"Opening document {file_path}")
                # Open the document
                doc = Document(file_path)
                
                # Copy content from the document
                # Skip the first few paragraphs (title, subtitle, etc.)
                skip_paragraphs = 3
                
                # Find the table in the document
                for table in doc.tables:
                    # Create a new table in the compiled document
                    new_table = compiled_doc.add_table(rows=1, cols=3)
                    new_table.style = 'Table Grid'
                    
                    # Set column widths
                    new_table.autofit = False
                    new_table.allow_autofit = False
                    
                    # Calculate available width
                    available_width = section.page_width - section.left_margin - section.right_margin
                    
                    # Set column widths
                    new_table.columns[0].width = int(available_width * 0.4)
                    new_table.columns[1].width = int(available_width * 0.2)
                    new_table.columns[2].width = int(available_width * 0.4)
                    
                    # Copy headers
                    header_cells = new_table.rows[0].cells
                    source_header_cells = table.rows[0].cells
                    
                    for j in range(min(len(header_cells), len(source_header_cells))):
                        header_cells[j].text = source_header_cells[j].text
                    
                    # Style headers
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(14)
                    
                    # Copy rows
                    for row_idx in range(1, len(table.rows)):
                        source_row = table.rows[row_idx]
                        new_row = new_table.add_row()
                        
                        # Copy cells
                        for cell_idx in range(min(len(source_row.cells), len(new_row.cells))):
                            source_cell = source_row.cells[cell_idx]
                            new_cell = new_row.cells[cell_idx]
                            
                            # Copy text
                            new_cell.text = source_cell.text
                            
                            # Style text
                            for paragraph in new_cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                for run in paragraph.runs:
                                    run.font.size = Pt(12)
                    
                    # Remove borders from table
                    for row in new_table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.space_after = Pt(12)
                            cell.border.top = cell.border.bottom = cell.border.left = cell.border.right = None
            except Exception as e:
                logger.error(f"Error processing document {file_path}: {str(e)}")
                logger.error(traceback.format_exc())
                error_para = compiled_doc.add_paragraph()
                error_para.add_run(f"Error processing document: {str(e)}").font.color.rgb = (255, 0, 0)
        
        # Save compiled document
        logger.info(f"Saving compiled document to {output_path}")
        try:
            compiled_doc.save(output_path)
            logger.info(f"Compiled document saved successfully to {output_path}")
            
            # Verify file exists after saving
            if os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                logger.info(f"Verified compiled file exists at {output_path} with size {file_size} bytes")
            else:
                logger.error(f"Compiled file does not exist after saving: {output_path}")
                
            return True
        except Exception as e:
            logger.error(f"Error saving compiled document to {output_path}: {str(e)}")
            logger.error(traceback.format_exc())
            return False
    except Exception as e:
        logger.error(f"Error compiling documents: {str(e)}")
        logger.error(traceback.format_exc())
        return False

def process_documents_thread(task_id, file_paths):
    """Process documents in a separate thread"""
    try:
        logger.info(f"Starting processing thread for task {task_id}")
        logger.info(f"Number of files to process: {len(file_paths)}")
        
        # Update task status
        tasks[task_id]['status'] = 'processing'
        tasks[task_id]['progress'] = 10
        tasks[task_id]['message'] = 'Processing documents...'
        
        processed_files = []
        processed_files_paths = []
        
        # Process each file
        for i, file_path in enumerate(file_paths):
            try:
                logger.info(f"Processing file {i+1}/{len(file_paths)}: {file_path}")
                
                # Update task status
                progress = 10 + int(80 * (i / len(file_paths)))
                tasks[task_id]['progress'] = progress
                tasks[task_id]['message'] = f'Processing file {i+1} of {len(file_paths)}...'
                
                # Get original filename
                original_filename = os.path.basename(file_path)
                name_without_ext = os.path.splitext(original_filename)[0]
                
                # Extract text from document
                logger.info(f"Extracting text from {file_path}")
                doc = Document(file_path)
                latin_text = ""
                
                for para in doc.paragraphs:
                    latin_text += para.text + "\n"
                
                logger.info(f"Extracted {len(latin_text)} characters of text")
                
                # Update task status
                tasks[task_id]['message'] = f'Correcting Latin text for {original_filename}...'
                
                # Correct Latin text
                corrected_latin = correct_latin_with_chatgpt(latin_text)
                
                # Update task status
                tasks[task_id]['message'] = f'Translating to Dutch for {original_filename}...'
                
                # Translate to Dutch
                dutch_translation = translate_latin_to_dutch_with_chatgpt(corrected_latin)
                
                # Update task status
                tasks[task_id]['message'] = f'Creating document for {original_filename}...'
                
                # Create output filename
                output_filename = f"processed_{name_without_ext}_{int(time.time())}.docx"
                output_path = os.path.join(PROCESSED_FOLDER, output_filename)
                
                logger.info(f"Creating document at {output_path}")
                
                # Create document
                success = create_three_column_document(corrected_latin, dutch_translation, output_path)
                
                if success:
                    logger.info(f"Document created successfully at {output_path}")
                    
                    # Add to processed files
                    processed_files.append({
                        'original_name': original_filename,
                        'processed_name': output_filename,
                        'download_url': f'/download/{output_filename}'
                    })
                    
                    # Add to processed files list for compilation
                    processed_files_paths.append(output_path)
                else:
                    logger.error(f"Failed to create document at {output_path}")
                    processed_files.append({
                        'original_name': original_filename,
                        'error': "Failed to create document"
                    })
            except Exception as e:
                logger.error(f"Error processing file {file_path}: {str(e)}")
                logger.error(traceback.format_exc())
                # Add error to processed files
                processed_files.append({
                    'original_name': os.path.basename(file_path),
                    'error': str(e)
                })
        
        # Update task status
        tasks[task_id]['progress'] = 90
        tasks[task_id]['message'] = 'Compiling documents...'
        
        # Compile documents if there are multiple files
        compiled_doc = None
        if len(processed_files_paths) > 1:
            logger.info("Compiling multiple documents")
            compiled_filename = f"compiled_{int(time.time())}.docx"
            compiled_path = os.path.join(PROCESSED_FOLDER, compiled_filename)
            
            if compile_documents(processed_files_paths, compiled_path):
                logger.info(f"Compilation successful: {compiled_path}")
                compiled_doc = {
                    'name': compiled_filename,
                    'download_url': f'/download/{compiled_filename}'
                }
            else:
                logger.error(f"Compilation failed: {compiled_path}")
        
        # Update task status
        tasks[task_id]['status'] = 'completed'
        tasks[task_id]['progress'] = 100
        tasks[task_id]['message'] = 'Processing completed'
        tasks[task_id]['processed_files'] = processed_files
        tasks[task_id]['compiled_doc'] = compiled_doc
        
        logger.info(f"Task {task_id} completed successfully")
    except Exception as e:
        logger.error(f"Error in processing thread for task {task_id}: {str(e)}")
        logger.error(traceback.format_exc())
        tasks[task_id]['status'] = 'error'
        tasks[task_id]['message'] = f'Error: {str(e)}'

# Routes
@app.route('/')
def index():
    logger.info("Serving index page")
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    logger.info("Received upload request")
    # Check if files were uploaded
    if 'files[]' not in request.files:
        logger.warning("No files uploaded")
        return jsonify({'error': 'No files uploaded'}), 400
    
    files = request.files.getlist('files[]')
    logger.info(f"Number of files uploaded: {len(files)}")
    
    # Check if files are empty
    if not files or files[0].filename == '':
        logger.warning("No files selected")
        return jsonify({'error': 'No files selected'}), 400
    
    # Create task
    task_id = str(uuid.uuid4())
    logger.info(f"Created task {task_id}")
    tasks[task_id] = {
        'status': 'uploaded',
        'progress': 0,
        'message': 'Files uploaded',
        'file_paths': []
    }
    
    # Process each file
    for file in files:
        # Check if file is allowed
        if file and allowed_file(file.filename):
            # Secure filename
            filename = secure_filename(file.filename)
            logger.info(f"Processing file: {filename}")
            
            # Add timestamp to filename to avoid collisions
            timestamp = int(time.time())
            filename_with_timestamp = f"{timestamp}_{filename}"
            
            # Save file
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename_with_timestamp)
            logger.info(f"Saving file to {file_path}")
            try:
                file.save(file_path)
                logger.info(f"File saved successfully to {file_path}")
                
                # Verify file exists after saving
                if os.path.exists(file_path):
                    file_size = os.path.getsize(file_path)
                    logger.info(f"Verified file exists at {file_path} with size {file_size} bytes")
                else:
                    logger.error(f"File does not exist after saving: {file_path}")
                
                # Add to task
                tasks[task_id]['file_paths'].append(file_path)
            except Exception as e:
                logger.error(f"Error saving file to {file_path}: {str(e)}")
                logger.error(traceback.format_exc())
        else:
            logger.warning(f"Invalid file: {file.filename}")
    
    # Check if any files were saved
    if not tasks[task_id]['file_paths']:
        logger.warning("No valid files uploaded")
        del tasks[task_id]
        return jsonify({'error': 'No valid files uploaded'}), 400
    
    logger.info(f"Upload successful for task {task_id}")
    return jsonify({'task_id': task_id}), 200

@app.route('/process/<task_id>', methods=['POST'])
def process_files(task_id):
    logger.info(f"Received process request for task {task_id}")
    # Check if task exists
    if task_id not in tasks:
        logger.warning(f"Task not found: {task_id}")
        return jsonify({'error': 'Task not found'}), 404
    
    # Check if task is already processing
    if tasks[task_id]['status'] == 'processing':
        logger.warning(f"Task {task_id} is already processing")
        return jsonify({'error': 'Task is already processing'}), 400
    
    logger.info(f"Starting processing for task {task_id}")
    
    # Start processing thread
    thread = threading.Thread(
        target=process_documents_thread,
        args=(task_id, tasks[task_id]['file_paths'])
    )
    thread.daemon = True
    thread.start()
    
    logger.info(f"Processing thread started for task {task_id}")
    return jsonify({'status': 'processing_started'}), 200

@app.route('/status/<task_id>')
def get_status(task_id):
    logger.info(f"Received status request for task {task_id}")
    # Check if task exists
    if task_id not in tasks:
        logger.warning(f"Task not found: {task_id}")
        return jsonify({'error': 'Task not found'}), 404
    
    # Return task status
    logger.info(f"Returning status for task {task_id}: {tasks[task_id]['status']}")
    return jsonify(tasks[task_id]), 200

@app.route('/download/<filename>')
def download_file(filename):
    logger.info(f"Received download request for file: {filename}")
    
    # Check if file exists
    file_path = os.path.join(app.config['PROCESSED_FOLDER'], filename)
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return jsonify({'error': 'File not found'}), 404
    
    logger.info(f"Sending file: {file_path}")
    try:
        return send_from_directory(app.config['PROCESSED_FOLDER'], filename, as_attachment=True)
    except Exception as e:
        logger.error(f"Error sending file {filename}: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'Error sending file: {str(e)}'}), 500

@app.route('/preview/<filename>')
def preview_file(filename):
    logger.info(f"Received preview request for file: {filename}")
    file_path = os.path.join(app.config['PROCESSED_FOLDER'], filename)
    
    # Check if file exists
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        abort(404)
    
    # For DOCX files, we can't preview directly in the browser
    # Return a placeholder or convert to PDF in a production environment
    if filename.endswith('.docx'):
        logger.info(f"Sending DOCX file for preview: {filename}")
        try:
            return send_from_directory(app.config['PROCESSED_FOLDER'], filename, as_attachment=False)
        except Exception as e:
            logger.error(f"Error sending file for preview {filename}: {str(e)}")
            logger.error(traceback.format_exc())
            return jsonify({'error': f'Error sending file for preview: {str(e)}'}), 500
    
    # For other files, return them directly
    logger.info(f"Sending file for preview: {filename}")
    try:
        return send_from_directory(app.config['PROCESSED_FOLDER'], filename)
    except Exception as e:
        logger.error(f"Error sending file for preview {filename}: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'Error sending file for preview: {str(e)}'}), 500

# Error handlers
@app.errorhandler(404)
def not_found(error):
    logger.warning(f"404 error: {request.path}")
    return jsonify({'error': 'Not found'}), 404

@app.errorhandler(413)
def request_entity_too_large(error):
    logger.warning(f"413 error: File too large")
    return jsonify({'error': 'File too large'}), 413

@app.errorhandler(500)
def internal_server_error(error):
    logger.error(f"500 error: {str(error)}")
    return jsonify({'error': 'Internal server error'}), 500

@app.route('/debug/log')
def view_log():
    """View the last 100 lines of the log file (for debugging)"""
    try:
        with open(log_file, 'r') as f:
            lines = f.readlines()
            last_lines = lines[-100:] if len(lines) > 100 else lines
            return Response(''.join(last_lines), mimetype='text/plain')
    except Exception as e:
        logger.error(f"Error reading log file: {str(e)}")
        return jsonify({'error': f'Error reading log file: {str(e)}'}), 500

@app.route('/debug/env')
def view_env():
    """View environment variables (for debugging)"""
    env_vars = {}
    for key in ['RENDER', 'RENDER_PERSISTENT_DIR', 'OPENAI_API_KEY']:
        if key in os.environ:
            value = os.environ[key]
            if key == 'OPENAI_API_KEY':
                value = value[:4] + '****' if value else 'Not set'
            env_vars[key] = value
        else:
            env_vars[key] = 'Not set'
    
    env_vars['UPLOAD_FOLDER'] = UPLOAD_FOLDER
    env_vars['PROCESSED_FOLDER'] = PROCESSED_FOLDER
    env_vars['IS_RENDER'] = IS_RENDER
    
    return jsonify(env_vars)

@app.route('/debug/files')
def view_files():
    """View files in upload and processed folders (for debugging)"""
    try:
        upload_files = os.listdir(UPLOAD_FOLDER) if os.path.exists(UPLOAD_FOLDER) else []
        processed_files = os.listdir(PROCESSED_FOLDER) if os.path.exists(PROCESSED_FOLDER) else []
        
        return jsonify({
            'upload_folder': UPLOAD_FOLDER,
            'upload_files': upload_files,
            'processed_folder': PROCESSED_FOLDER,
            'processed_files': processed_files
        })
    except Exception as e:
        logger.error(f"Error listing files: {str(e)}")
        return jsonify({'error': f'Error listing files: {str(e)}'}), 500

if __name__ == '__main__':
    # Get port from environment variable or use default
    port = int(os.environ.get('PORT', 5000))
    
    logger.info(f"Starting server on port {port}")
    
    # Run the app
    app.run(host='0.0.0.0', port=port, debug=False)
